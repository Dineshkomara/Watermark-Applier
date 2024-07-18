import cv2
from flask import Flask, render_template, request, send_file
from PIL import ImageFont
import PIL.ImageDraw
import PIL.Image
import io
from rembg import remove
from werkzeug.utils import secure_filename
import os
from werkzeug.utils import secure_filename
from docx import Document
from spire.doc import *
from spire.doc.common import *
from spire.pdf import *
from spire.pdf.common import *
import math

app = Flask(__name__)

# Function to apply text watermark to an image
def apply_text_watermark(input_image, watermark_text, x_position,y_position, transparency, text_color, font_name, font_size):
    image = PIL.Image.open(input_image)
    if image.mode == 'RGBA':
        image = image.convert('RGB')  # Convert RGBA to RGB
    draw = PIL.ImageDraw.Draw(image)
    font1 = ImageFont.truetype(font_name, font_size)  # You can change the font and size as needed
    text_width = draw.textlength(watermark_text, font1)
    text_height = font_size
    x = int((image.width - text_width) * x_position / 100)
    y = int((image.height - text_height) * y_position / 100)
    text_color = (*text_color[:3], int(255 * transparency))
    draw.text((x, y), watermark_text, fill=text_color, font=font1)
    return image

# Function to apply image watermark to an image
# Function to apply image watermark to an image
def apply_image_watermark(input_image, watermark_image, x_position,y_position, transparency, remove_background,scale=100):
    image = PIL.Image.open(input_image)
    watermark = PIL.Image.open(watermark_image)
    watermark_width, watermark_height = watermark.size
    watermark_width,watermark_height=watermark_width*float(float(scale)/100),watermark_height*float(float(scale)/100)
    x = int((image.width - watermark_width) * x_position / 100)
    y = int((image.height - watermark_height) * y_position / 100)
    
    # Apply background removal if specified
    if remove_background == 'yes':
        watermark = remove_background_watermark(watermark)
    
    # Apply transparency to the watermark
    watermark = apply_transparency(watermark, transparency)
    # Paste the watermark onto the input image
    image.paste(watermark, (x, y), watermark)
    return image
def apply_text_watermark_to_pdf(input_pdf, watermark_text, font_name="Helvetica", font_size=36, transparency=0.5):
    input_pdf.save("input_file.pdf")
    pdf = PdfDocument()
    pdf.LoadFromFile("input_file.pdf")
    font = PdfTrueTypeFont(font_name, float(font_size), 0, True)
    text = watermark_text
    set1 = float (font.MeasureString(text).Width * math.sqrt(2) / 4)
    set2 = float (font.MeasureString(text).Height * math.sqrt(2) / 4)
    for i in range(pdf.Pages.Count):
        page = pdf.Pages.get_Item(i)
        page.Canvas.SetTransparency(float(transparency))
        page.Canvas.TranslateTransform((page.Canvas.Size.Width / 2 - set1 - set2),page.Canvas.Size.Height / 2 + set1 - set2)
        page.Canvas.RotateTransform(-45.0)
        page.Canvas.DrawString(text, font, PdfBrushes.get_Black(), 0.0, 0.0)
    pdf.SaveToFile("output/SingleLineTextWatermark.pdf")
    pdf.Close()

def apply_image_watermark_to_pdf(input_pdf, watermark_image, transparency=0.5 ,remove_background="no",scale1=100):
    input_pdf.save("input_file.pdf")
    watermark=PIL.Image.open(watermark_image)
    if remove_background == 'yes':
        watermark = remove_background_watermark(watermark)
    watermark = apply_transparency(watermark, transparency)
    watermark.save("watermark.png")
    pdf = PdfDocument()
    pdf.LoadFromFile("input_file.pdf")
    image = PdfImage.FromFile("watermark.png")
    imageWidth = float(image.Width)*float(float(scale1)/100)
    imageHeight = float(image.Height)*float(float(scale1)/100)
    for i in range(pdf.Pages.Count):
        page = pdf.Pages.get_Item(i)
        page.Canvas.SetTransparency(transparency)
        pageWidth = page.ActualSize.Width
        pageHeight = page.ActualSize.Height
        page.Canvas.DrawImage(image, pageWidth/2 - imageWidth/2, pageHeight/2 - imageHeight/2, imageWidth, imageHeight)
    pdf.SaveToFile("output/SingleImageWatermark.pdf")
    pdf.Close()
 

def apply_text_watermark_to_docx(docx_file, watermark_text, font_name='Calibri', font_size=40, font_color=[0,0,0], transparency=0.5):
    docx_file.save("input_file.docx")
    document = Document()
    document.LoadFromFile("input_file.docx")
    txtWatermark = TextWatermark()
    # Set the format of the text watermark
    txtWatermark.Text = watermark_text
    txtWatermark.FontName = font_name
    txtWatermark.FontSize = font_size
    txtWatermark.Color.FromArgb(int(255*transparency),font_color[0],font_color[1],font_color[2])
    txtWatermark.Layout = WatermarkLayout.Diagonal
    document.Watermark = txtWatermark
    #Save the result document
    document.SaveToFile("Output/TextWatermark.docx")
    document.Close()

def apply_image_watermark_to_docx(input_docx, watermark_image, transparency,remove_background,scale=100):
    input_docx.save("input_file.docx")
    document = Document()
    document.LoadFromFile("input_file.docx")
    picture = PictureWatermark()
    watermark=PIL.Image.open(watermark_image)
    if remove_background == 'yes':
        watermark = remove_background_watermark(watermark)
    # Apply transparency to the watermark
    watermark = apply_transparency(watermark, transparency)
    watermark.save("logo.png")
    picture.SetPicture("logo.png")
    picture.Scaling = int(scale)
    picture.IsWashout = False
    document.Watermark = picture
    document.SaveToFile("Output/ImageWatermark.docx")
    document.Close()
    doc=Document()
    doc.LoadFromFile("Output/ImageWatermark.docx")
    return doc

# Function to apply transparency to an image
def apply_transparency(image, transparency):
    if image.mode != 'RGBA':
        image = image.convert('RGBA')
    
    alpha = image.split()[3]  # Get the alpha channel
    alpha = alpha.point(lambda p: p * transparency)  # Apply transparency
    image.putalpha(alpha)  # Update the alpha channel
    
    return image


# Function to remove background from the watermark image
def remove_background_watermark(watermark):
    watermark=remove(watermark)
    return watermark



@app.route('/')
def index():
    return render_template('index.html')

@app.route('/process', methods=['POST'])
def process():
    file = request.files['file']
    file_option=request.form.get('file_option')
    watermark_option = request.form.get('watermark_option')
    x_position = int(request.form.get('x_position', 50))  # Retrieve X position
    y_position = int(request.form.get('y_position', 50))
    transparency = float(request.form.get('transparency', 0.5))
    remove_background = request.form.get('remove_background', 'yes')
    scale=request.form.get('scale',100)

    if file.filename == '':
        return "No file uploaded", 400
    filename = secure_filename(file.filename)
    file_ext = os.path.splitext(filename)[1]
    
    if file_ext.lower() in ['.jpg', '.jpeg', '.png', '.gif']:
        if watermark_option == 'text':
            watermark_text = request.form.get('watermark_text')
            text_color_hex = request.form.get('text_color', '#000000')
            text_color = tuple(int(text_color_hex[i:i+2], 16) for i in (1, 3, 5))
            font_name = request.form.get('font_name', 'arial.ttf')
            font_size = int(request.form.get('font_size', 50))
            if not watermark_text:
                return "Watermark text not provided", 400
            watermarked_image = apply_text_watermark(file, watermark_text, x_position, y_position, transparency, text_color, font_name, font_size)
        elif watermark_option == 'image':
            watermark_image = request.files.get('watermark_image')
            if not watermark_image:
                return "No image selected for watermark", 400
            watermarked_image = apply_image_watermark(file, watermark_image, x_position, y_position, transparency, remove_background)

        output = io.BytesIO()
        watermarked_image.save("output.png")
        output.seek(0)
        return send_file("output.png", mimetype='image/png')
    elif file_ext.lower() == '.pdf':
        input_file = file
        if watermark_option == 'text':
            watermark_text = request.form.get('watermark_text')
            text_color_hex = request.form.get('text_color', '#000000')
            text_color = tuple(int(text_color_hex[i:i+2], 16) for i in (1, 3, 5))
            font_name = request.form.get('font_name', 'Helvetica')
            font_size = int(request.form.get('font_size', 36))
            if not watermark_text:
                return "Watermark text not provided", 400
            apply_text_watermark_to_pdf(input_file, watermark_text, font_name, font_size, transparency)
            return send_file("output/SingleLineTextWatermark.pdf", mimetype='application/pdf')
        elif watermark_option == 'image':
            watermark_image = request.files.get('watermark_image')
            if not watermark_image:
                return "No image selected for watermark", 400
            apply_image_watermark_to_pdf(input_file, watermark_image, transparency,remove_background,scale)
            return send_file("output/SingleImageWatermark.pdf", mimetype='application/pdf')
    elif file_ext.lower() == '.docx':
        input_file = file
        if watermark_option == 'text':
            watermark_text = request.form.get('watermark_text')
            text_color_hex = request.form.get('text_color', '#000000')
            temp = tuple(int(text_color_hex[i:i+2], 16) for i in (1, 3, 5))
            text_color=list(temp)
            font_name = request.form.get('font_name', 'Helvetica')
            font_size = int(request.form.get('font_size', 36))
            apply_text_watermark_to_docx(input_file, watermark_text, font_name, font_size, text_color, transparency)
            return send_file("Output/TextWatermark.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        elif watermark_option == 'image':
            watermark_image = request.files.get('watermark_image')
            if not watermark_image:
                return "No image selected for watermark", 400
            apply_image_watermark_to_docx(input_file, watermark_image, transparency, remove_background,scale)
            return send_file("Output/ImageWatermark.docx", mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    else:
        return f"Unsupported file format", 400

if __name__ == '__main__':
    app.run(debug=True)